"""
Document parsing service for PDF and DOCX files with OCR support.

Extracts structured text, section hierarchy, and image text from academic
documents. Returns a ParsedDocument with full_text, sections, images_text,
and rich metadata (page_count, language, word_count, reading_time, etc.).
"""
from __future__ import annotations

import io
import logging
import re
from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional, Tuple

import fitz  # PyMuPDF
import pytesseract
from docx import Document as DocxDocument
from docx.oxml.ns import qn
from PIL import Image

from app.config import settings

logger = logging.getLogger(__name__)

# Optional language detection — gracefully absent if not installed
try:
    from langdetect import detect as _langdetect_fn
    from langdetect.lang_detect_exception import LangDetectException
    _HAS_LANGDETECT = True
except ImportError:
    _HAS_LANGDETECT = False


# ---------------------------------------------------------------------------
# Data structures
# ---------------------------------------------------------------------------

@dataclass
class ParsedSection:
    """A logical section of text extracted from a document."""

    title: str          # Heading text; empty string for body-only sections
    content: str        # Body text belonging to this section
    level: int          # Heading depth: 0 = no heading, 1 = H1, 2 = H2, 3 = H3+
    page_num: int = 0   # Starting page number (1-based; 0 = unknown)


@dataclass
class ParsedDocument:
    """
    Output of the DocumentParser.

    Attributes:
        full_text:    Complete text of the document (joined from all pages/sections).
        sections:     Ordered list of ParsedSection objects.
        images_text:  OCR text extracted from embedded images/diagrams.
        metadata:     Dict with keys: page_count, detected_language, has_images,
                      word_count, reading_time_minutes, title, author, subject,
                      file_type, and any format-specific fields.
    """

    full_text: str
    sections: List[ParsedSection] = field(default_factory=list)
    images_text: List[str] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)


# ---------------------------------------------------------------------------
# Parser
# ---------------------------------------------------------------------------

class DocumentParser:
    """Parses PDF and DOCX documents into structured ParsedDocument objects."""

    def __init__(self) -> None:
        if settings.TESSERACT_CMD:
            pytesseract.pytesseract.tesseract_cmd = settings.TESSERACT_CMD

    async def parse_document(self, file_path: str, file_type: str) -> ParsedDocument:
        """
        Parse a document file and return a structured ParsedDocument.

        Args:
            file_path: Absolute path to the file on disk.
            file_type: Extension with or without dot, e.g. ".pdf" or "docx".

        Returns:
            ParsedDocument with text, sections, image OCR, and metadata.

        Raises:
            ValueError:   Unsupported file type.
            RuntimeError: Password-protected or unreadable file.
        """
        ft = file_type.lower().lstrip(".")
        if ft == "pdf":
            return await self._parse_pdf(file_path)
        elif ft in ("docx", "doc"):
            return await self._parse_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type!r}")

    # ------------------------------------------------------------------
    # PDF
    # ------------------------------------------------------------------

    async def _parse_pdf(self, file_path: str) -> ParsedDocument:
        """Parse a PDF using PyMuPDF with OCR fallback for image-only pages."""
        try:
            doc = fitz.open(file_path)
        except Exception as exc:
            raise RuntimeError(f"Cannot open PDF file: {exc}") from exc

        if doc.needs_pass:
            doc.close()
            raise RuntimeError(
                "PDF is password-protected. Please provide an unlocked copy."
            )

        raw_meta = doc.metadata or {}
        page_count = doc.page_count

        # ---- Pass 1: collect all span font sizes to find the body size ----
        font_sizes: List[float] = []
        for page in doc:
            for block in page.get_text("dict")["blocks"]:
                if block.get("type") != 0:
                    continue
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        sz = span.get("size", 0.0)
                        if sz > 0:
                            font_sizes.append(sz)

        body_font_size = _modal_font_size(font_sizes) if font_sizes else 11.0
        # A span is a heading candidate if its font is ≥ 15 % larger than body
        heading_size_threshold = body_font_size * 1.15

        # ---- Pass 2: extract text, sections, and images ----
        sections: List[ParsedSection] = []
        images_text: List[str] = []
        has_images = False
        all_page_texts: List[str] = []

        # Start with an anonymous "preamble" section that captures any text
        # appearing before the first detected heading.
        current_section = ParsedSection(title="", content="", level=0, page_num=1)

        for page_num, page in enumerate(doc, start=1):
            page_height = page.rect.height

            # Zones to discard: top 8 % (running header) and bottom 8 % (footer)
            header_cutoff = page_height * 0.08
            footer_cutoff = page_height * 0.92

            raw_blocks = page.get_text("dict")["blocks"]
            page_has_text = False

            # Collect (y0, x0, text, is_heading, span_size) for this page
            line_items: List[Tuple[float, float, str, bool, float]] = []

            for block in raw_blocks:
                btype = block.get("type", -1)
                bbox = block.get("bbox", [0, 0, 0, 0])
                y0, x0 = bbox[1], bbox[0]

                if btype == 1:
                    # Image block — run OCR
                    has_images = True
                    ocr = await self._ocr_block(page, block)
                    if ocr.strip():
                        images_text.append(ocr.strip())
                    continue

                if btype != 0:
                    continue

                # Discard header/footer bands
                if y0 < header_cutoff or y0 > footer_cutoff:
                    continue

                for line in block.get("lines", []):
                    max_sz = 0.0
                    is_bold_line = False
                    span_parts: List[str] = []

                    for span in line.get("spans", []):
                        raw_txt = span.get("text", "")
                        if not raw_txt.strip():
                            continue
                        sz = span.get("size", 0.0)
                        flags = span.get("flags", 0)
                        if sz > max_sz:
                            max_sz = sz
                        if flags & 16:          # bit 4 = bold
                            is_bold_line = True
                        span_parts.append(raw_txt)

                    line_text = " ".join(span_parts).strip()

                    # Skip isolated page numbers (pure digits ≤ 4 chars)
                    if re.match(r"^\d{1,4}$", line_text):
                        continue

                    if not line_text:
                        continue

                    page_has_text = True

                    # Heading heuristic: significantly larger font OR
                    # bold + body-size font + short text (≤ 15 words)
                    word_count = len(line_text.split())
                    is_heading = max_sz >= heading_size_threshold or (
                        is_bold_line
                        and max_sz >= body_font_size
                        and word_count <= 15
                    )

                    line_items.append((y0, x0, line_text, is_heading, max_sz))

            if not page_has_text:
                # Image-only page: full-page OCR
                ocr = await self._ocr_page(page)
                if ocr.strip():
                    current_section.content += "\n" + ocr
                    all_page_texts.append(f"[Page {page_num}]\n{ocr}")
                continue

            # Sort by vertical position (primary) then horizontal (secondary)
            # so multi-column text reads left-column before right-column
            line_items.sort(key=lambda item: (item[0], item[1]))

            page_lines: List[str] = []
            for _y, _x, line_text, is_heading, span_sz in line_items:
                if is_heading:
                    # Flush existing section
                    if current_section.title or current_section.content.strip():
                        sections.append(current_section)
                    level = _estimate_heading_level(span_sz, body_font_size)
                    current_section = ParsedSection(
                        title=line_text,
                        content="",
                        level=level,
                        page_num=page_num,
                    )
                    page_lines.append(f"\n{'#' * level} {line_text}")
                else:
                    current_section.content += " " + line_text
                    page_lines.append(line_text)

            all_page_texts.append(f"[Page {page_num}]\n" + "\n".join(page_lines))

        # Flush last section
        if current_section.title or current_section.content.strip():
            sections.append(current_section)

        doc.close()

        # Also extract tables using find_tables() if available (PyMuPDF ≥ 1.23)
        table_texts: List[str] = []
        try:
            pdf2 = fitz.open(file_path)
            for page in pdf2:
                tables = page.find_tables()
                for table in tables:
                    rows = table.extract()
                    if rows:
                        formatted = _format_table_rows(rows)
                        if formatted:
                            table_texts.append(formatted)
            pdf2.close()
        except (AttributeError, Exception):
            pass  # find_tables not available or failed — non-critical

        full_text = "\n\n".join(all_page_texts)
        if table_texts:
            full_text += "\n\n[Tables]\n" + "\n\n".join(table_texts)

        word_count = len(full_text.split())

        metadata: Dict[str, Any] = {
            "page_count": page_count,
            "detected_language": _detect_language(full_text[:3000]),
            "has_images": has_images,
            "word_count": word_count,
            "reading_time_minutes": round(word_count / 200, 1),
            "title": raw_meta.get("title", ""),
            "author": raw_meta.get("author", ""),
            "subject": raw_meta.get("subject", ""),
            "creator": raw_meta.get("creator", ""),
            "creation_date": raw_meta.get("creationDate", ""),
            "file_type": "pdf",
        }

        return ParsedDocument(
            full_text=full_text,
            sections=sections,
            images_text=images_text,
            metadata=metadata,
        )

    async def _ocr_page(self, page: fitz.Page) -> str:
        """Render an entire page at 2× scale and run Tesseract OCR."""
        try:
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            return pytesseract.image_to_string(img)
        except Exception as exc:
            logger.warning(f"Full-page OCR failed on page {page.number + 1}: {exc}")
            return ""

    async def _ocr_block(self, page: fitz.Page, block: dict) -> str:
        """Render a single image block and run Tesseract OCR."""
        try:
            bbox = fitz.Rect(block["bbox"])
            # Add a small padding so text at the very edge isn't clipped
            clip = bbox + fitz.Rect(-4, -4, 4, 4)
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), clip=clip)
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            return pytesseract.image_to_string(img)
        except Exception as exc:
            logger.warning(f"Image block OCR failed: {exc}")
            return ""

    # ------------------------------------------------------------------
    # DOCX
    # ------------------------------------------------------------------

    async def _parse_docx(self, file_path: str) -> ParsedDocument:
        """Parse a DOCX file preserving heading hierarchy and tables."""
        try:
            doc = DocxDocument(file_path)
        except Exception as exc:
            raise RuntimeError(f"Cannot open DOCX file: {exc}") from exc

        HEADING_STYLES: Dict[str, int] = {
            "heading 1": 1,
            "heading 2": 2,
            "heading 3": 3,
            "heading 4": 3,
            "heading 5": 3,
            "title": 1,
            "subtitle": 2,
        }

        # XML tag names for embedded images
        BLIP_TAG = qn("a:blip")
        EMBED_KEY = qn("r:embed")

        sections: List[ParsedSection] = []
        images_text: List[str] = []
        has_images = False
        all_text_parts: List[str] = []

        current_section = ParsedSection(title="", content="", level=0, page_num=0)

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style_name = (
                para.style.name.lower() if para.style and para.style.name else ""
            )
            heading_level = HEADING_STYLES.get(style_name, 0)

            # Treat a short, entirely-bold paragraph as an implicit H3
            if heading_level == 0 and _is_implicit_heading(para):
                heading_level = 3

            if heading_level > 0:
                # Flush previous section
                if current_section.title or current_section.content.strip():
                    sections.append(current_section)
                current_section = ParsedSection(
                    title=text,
                    content="",
                    level=heading_level,
                    page_num=0,
                )
                all_text_parts.append(f"\n{'#' * heading_level} {text}\n")
            else:
                current_section.content += " " + text
                all_text_parts.append(text)

            # Extract images embedded in this paragraph
            for blip in para._element.iter(BLIP_TAG):
                rid = blip.get(EMBED_KEY)
                if not rid:
                    continue
                has_images = True
                try:
                    img_part = doc.part.related_parts[rid]
                    img = Image.open(io.BytesIO(img_part.blob))
                    ocr = pytesseract.image_to_string(img)
                    if ocr.strip():
                        images_text.append(ocr.strip())
                except Exception as exc:
                    logger.warning(f"DOCX image OCR failed: {exc}")

        # Tables
        for table in doc.tables:
            rows: List[str] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                non_empty = [c for c in cells if c]
                if non_empty:
                    rows.append(" | ".join(non_empty))
            if rows:
                table_text = "\n".join(rows)
                current_section.content += "\n" + table_text
                all_text_parts.append(table_text)

        # Flush final section
        if current_section.title or current_section.content.strip():
            sections.append(current_section)

        core = doc.core_properties
        full_text = "\n\n".join(all_text_parts)
        word_count = len(full_text.split())

        metadata: Dict[str, Any] = {
            "page_count": None,   # python-docx cannot report rendered page count
            "detected_language": _detect_language(full_text[:3000]),
            "has_images": has_images,
            "word_count": word_count,
            "reading_time_minutes": round(word_count / 200, 1),
            "title": core.title or "",
            "author": core.author or "",
            "subject": core.subject or "",
            "created": str(core.created) if core.created else "",
            "modified": str(core.modified) if core.modified else "",
            "file_type": "docx",
        }

        return ParsedDocument(
            full_text=full_text,
            sections=sections,
            images_text=images_text,
            metadata=metadata,
        )

    # ------------------------------------------------------------------
    # Legacy compatibility
    # ------------------------------------------------------------------

    async def extract_metadata_only(
        self, file_path: str, file_type: str
    ) -> Dict[str, Any]:
        """Return only document metadata without full text extraction."""
        ft = file_type.lower().lstrip(".")
        try:
            if ft == "pdf":
                doc = fitz.open(file_path)
                meta = {
                    "page_count": doc.page_count,
                    "title": doc.metadata.get("title", ""),
                    "author": doc.metadata.get("author", ""),
                    "subject": doc.metadata.get("subject", ""),
                }
                doc.close()
                return meta
            elif ft in ("docx", "doc"):
                doc = DocxDocument(file_path)
                core = doc.core_properties
                return {
                    "title": core.title or "",
                    "author": core.author or "",
                    "subject": core.subject or "",
                }
        except Exception as exc:
            logger.error(f"Error extracting metadata from {file_path!r}: {exc}")
        return {}


# ---------------------------------------------------------------------------
# Module-level helpers
# ---------------------------------------------------------------------------

def _modal_font_size(sizes: List[float]) -> float:
    """Return the most frequently occurring font size (proxy for body text)."""
    rounded = [round(s, 1) for s in sizes]
    freq: Dict[float, int] = {}
    for s in rounded:
        freq[s] = freq.get(s, 0) + 1
    return max(freq, key=lambda k: freq[k])


def _estimate_heading_level(span_size: float, body_size: float) -> int:
    """Map a span's font-size ratio to an H1/H2/H3 level."""
    ratio = span_size / body_size if body_size > 0 else 1.0
    if ratio >= 1.5:
        return 1
    if ratio >= 1.25:
        return 2
    return 3


def _is_implicit_heading(para) -> bool:
    """Return True if a DOCX paragraph looks like an unlabelled heading.

    Criteria: short text (≤ 15 words) where every non-whitespace run is bold.
    """
    text = para.text.strip()
    if not text or len(text.split()) > 15:
        return False
    runs_with_text = [r for r in para.runs if r.text.strip()]
    return bool(runs_with_text) and all(r.bold for r in runs_with_text)


def _detect_language(sample: str) -> str:
    """Detect the language of a text sample; returns an ISO 639-1 code or 'unknown'."""
    if not _HAS_LANGDETECT or len(sample.split()) < 20:
        return "unknown"
    try:
        return _langdetect_fn(sample)
    except Exception:
        return "unknown"


def _format_table_rows(rows: List[List[Optional[str]]]) -> str:
    """Format a list-of-lists table as pipe-delimited text."""
    lines: List[str] = []
    for row in rows:
        cells = [str(cell).strip() if cell is not None else "" for cell in row]
        non_empty = [c for c in cells if c]
        if non_empty:
            lines.append(" | ".join(cells))
    return "\n".join(lines)
