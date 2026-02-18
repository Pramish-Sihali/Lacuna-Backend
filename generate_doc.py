"""Generate a DOCX document explaining the Lacuna backend architecture."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# --- Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def add_body(text):
    p = doc.add_paragraph(text)
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(f" {text}")
    else:
        p.add_run(text)
    return p

def add_code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(40, 40, 40)
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    return p

def add_comparison(nextjs, fastapi):
    """Add a Next.js vs FastAPI comparison row."""
    p = doc.add_paragraph()
    run1 = p.add_run("Next.js: ")
    run1.bold = True
    run1.font.color.rgb = RGBColor(0, 100, 0)
    p.add_run(nextjs)
    p.add_run("  |  ")
    run2 = p.add_run("FastAPI: ")
    run2.bold = True
    run2.font.color.rgb = RGBColor(0, 0, 150)
    p.add_run(fastapi)

# ==================== DOCUMENT START ====================

# Title
title = doc.add_heading('Lacuna Backend - Architecture Deep Dive', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Prepared for Mentor Discussion')
run.italic = True
run.font.size = Pt(12)

doc.add_paragraph()

# ==================== SECTION 1: WHAT IS LACUNA ====================
add_heading('1. What is Lacuna?', level=1)
add_body(
    'Lacuna is an AI-powered research discovery platform that helps academics and researchers '
    'visualize knowledge gaps in their document collections. Users upload research papers (PDF/DOCX), '
    'and the system automatically extracts concepts, detects relationships between them, identifies '
    'knowledge gaps, and provides a RAG (Retrieval-Augmented Generation) interface to query across '
    'all their documents.'
)

add_heading('Core Problem We Solve:', level=2)
add_bullet('Researchers have dozens/hundreds of papers but struggle to see the "big picture"')
add_bullet('Hard to identify what topics are under-explored or contradictory across papers')
add_bullet('No easy way to query across multiple documents semantically')

add_heading('What The System Does (End-to-End):', level=2)
add_bullet('Upload documents (PDF/DOCX with OCR support)', bold_prefix='Ingest:')
add_bullet('LLM extracts concepts, claims, and relationships from text', bold_prefix='Extract:')
add_bullet('HDBSCAN clustering + hierarchy detection groups related concepts', bold_prefix='Organize:')
add_bullet('Algorithmic detection of missing links, contradictions, under-explored areas', bold_prefix='Detect Gaps:')
add_bullet('Semantic search + LLM reasoning across all documents', bold_prefix='Query (RAG):')
add_bullet('Aggregates agreement/disagreement across documents', bold_prefix='Consensus:')

# ==================== SECTION 2: FASTAPI vs NEXTJS ====================
add_heading('2. FastAPI Explained (For Someone Who Knows Next.js)', level=1)

add_body(
    'If you know Next.js, you already understand 70% of the concepts. FastAPI is Python\'s '
    'equivalent of building a backend API. Here\'s how the mental models map:'
)

add_heading('2.1 The Core Analogy', level=2)

# Table comparing Next.js to FastAPI
table = doc.add_table(rows=12, cols=2)
table.style = 'Light Grid Accent 1'
# Header
table.rows[0].cells[0].text = 'Next.js Concept'
table.rows[0].cells[1].text = 'FastAPI Equivalent'
for cell in table.rows[0].cells:
    for p in cell.paragraphs:
        p.runs[0].bold = True

comparisons = [
    ('pages/api/users.ts (API Route)', 'routers/documents.py (Router)'),
    ('export default handler(req, res)', '@router.post("/upload") async def upload()'),
    ('req.body (request data)', 'Pydantic schema (validated automatically)'),
    ('res.json({...}) (response)', 'return {...} (auto-serialized to JSON)'),
    ('middleware.ts', 'app.add_middleware(CORSMiddleware, ...)'),
    ('next.config.js', 'config.py + .env file'),
    ('prisma / drizzle (ORM)', 'SQLAlchemy (ORM)'),
    ('prisma migrate', 'Alembic (migrations)'),
    ('lib/ or utils/ folder', 'services/ folder'),
    ('TypeScript interfaces', 'Pydantic models (schemas.py)'),
    ('npm run dev', 'uvicorn app.main:app --reload'),
]
for i, (nxt, fapi) in enumerate(comparisons, 1):
    table.rows[i].cells[0].text = nxt
    table.rows[i].cells[1].text = fapi

add_heading('2.2 How a Request Flows (Side-by-Side)', level=2)

add_body('Next.js API Route:')
add_code(
    '// pages/api/documents.ts\n'
    'export default async function handler(req: NextRequest, res: NextResponse) {\n'
    '  if (req.method === "POST") {\n'
    '    const body = await req.json();\n'
    '    const doc = await prisma.document.create({ data: body });\n'
    '    return res.json(doc);\n'
    '  }\n'
    '}'
)

add_body('FastAPI Equivalent:')
add_code(
    '# routers/documents.py\n'
    '@router.post("/upload")\n'
    'async def upload_document(\n'
    '    file: UploadFile = File(...),      # <-- auto file handling\n'
    '    db: AsyncSession = Depends(get_db)  # <-- dependency injection (like context)\n'
    '):\n'
    '    # ... process file ...\n'
    '    db.add(document)\n'
    '    await db.commit()\n'
    '    return {"id": document.id, "filename": document.filename}'
)

add_heading('2.3 Key Difference: Dependency Injection', level=2)
add_body(
    'FastAPI\'s killer feature is Depends(). In Next.js, you manually create DB connections or '
    'import singletons. In FastAPI, you declare what you need and the framework injects it:'
)
add_code(
    '# This is like React\'s useContext() but for the backend\n'
    'async def get_db():\n'
    '    async with async_session() as session:\n'
    '        yield session  # provides DB session, auto-closes after request\n'
    '\n'
    '# Any endpoint can request a DB session:\n'
    '@router.get("/")\n'
    'async def list_docs(db: AsyncSession = Depends(get_db)):\n'
    '    # db is automatically provided and cleaned up'
)

add_heading('2.4 Auto-Generated API Docs', level=2)
add_body(
    'Unlike Next.js where you need Swagger/OpenAPI setup manually, FastAPI auto-generates '
    'interactive API docs at /docs (Swagger UI) and /redoc. Every endpoint, request body, '
    'and response schema is documented automatically from your Python type hints.'
)

# ==================== SECTION 3: MVC IN FASTAPI ====================
add_heading('3. How We Implement MVC (Our Architecture Pattern)', level=1)

add_body(
    'We use a variation of MVC adapted for FastAPI. Technically it\'s closer to a '
    '"Service-Repository" pattern, but here\'s how the MVC concepts map:'
)

table2 = doc.add_table(rows=5, cols=3)
table2.style = 'Light Grid Accent 1'
table2.rows[0].cells[0].text = 'MVC Layer'
table2.rows[0].cells[1].text = 'Our Implementation'
table2.rows[0].cells[2].text = 'Files'
for cell in table2.rows[0].cells:
    for p in cell.paragraphs:
        p.runs[0].bold = True

mvc_rows = [
    ('Model', 'SQLAlchemy ORM models + Pydantic schemas', 'models/database_models.py\nmodels/schemas.py'),
    ('View', 'Routers (API endpoints that return JSON)', 'routers/documents.py\nrouters/concepts.py\nrouters/brain.py'),
    ('Controller', 'Service classes (business logic)', 'services/*.py (9 service files)'),
    ('Database', 'Async PostgreSQL + pgvector', 'database.py + docker-compose.yml'),
]
for i, (layer, impl, files) in enumerate(mvc_rows, 1):
    table2.rows[i].cells[0].text = layer
    table2.rows[i].cells[1].text = impl
    table2.rows[i].cells[2].text = files

add_heading('3.1 The Flow: Router -> Service -> Model -> Database', level=2)
add_code(
    'User Request: POST /api/documents/upload\n'
    '       |\n'
    '       v\n'
    '[Router: documents.py]  <-- Handles HTTP, validates input\n'
    '       |\n'
    '       v\n'
    '[Service: document_parser.py]  <-- Business logic (parse PDF)\n'
    '[Service: chunking.py]         <-- Business logic (split text)\n'
    '[Service: embedding.py]        <-- Business logic (generate vectors)\n'
    '       |\n'
    '       v\n'
    '[Model: database_models.py]  <-- ORM objects (Document, Chunk)\n'
    '       |\n'
    '       v\n'
    '[Database: PostgreSQL + pgvector]  <-- Persistent storage'
)

# ==================== SECTION 4: PROJECT STRUCTURE ====================
add_heading('4. Complete Project Structure', level=1)

add_code(
    'Lacuna-backend/\n'
    '|\n'
    '+-- app/\n'
    '|   +-- main.py                 # App entry point, middleware, lifespan events\n'
    '|   +-- config.py               # All settings from .env (Pydantic Settings)\n'
    '|   +-- database.py             # Async DB engine, session factory, init\n'
    '|   |\n'
    '|   +-- models/\n'
    '|   |   +-- database_models.py  # 7 SQLAlchemy tables (ORM)\n'
    '|   |   +-- schemas.py          # Pydantic request/response models\n'
    '|   |\n'
    '|   +-- routers/\n'
    '|   |   +-- health.py           # GET /api/health\n'
    '|   |   +-- documents.py        # CRUD for documents + upload pipeline\n'
    '|   |   +-- concepts.py         # Extract, cluster, detect gaps, concept map\n'
    '|   |   +-- brain.py            # RAG queries + consensus building\n'
    '|   |\n'
    '|   +-- services/               # ALL business logic lives here\n'
    '|   |   +-- document_parser.py  # PDF/DOCX parsing with OCR fallback\n'
    '|   |   +-- chunking.py         # Text -> chunks (500 tokens, 50 overlap)\n'
    '|   |   +-- embedding.py        # Ollama embeddings (768-dim vectors)\n'
    '|   |   +-- llm_extractor.py    # LLM concept/claim/relationship extraction\n'
    '|   |   +-- normalizer.py       # Concept deduplication (name + embedding)\n'
    '|   |   +-- clustering.py       # HDBSCAN clustering + hierarchy\n'
    '|   |   +-- relationships.py    # Concept relationship detection + scoring\n'
    '|   |   +-- gap_detector.py     # 4 types of knowledge gap detection\n'
    '|   |   +-- brain_service.py    # RAG pipeline + consensus builder\n'
    '|   |\n'
    '|   +-- utils/\n'
    '|       +-- helpers.py          # cosine_similarity, text cleaning, hashing\n'
    '|\n'
    '+-- alembic/                    # Database migration scripts\n'
    '+-- uploads/                    # Stored uploaded files\n'
    '+-- docker-compose.yml          # PostgreSQL 16 + pgvector container\n'
    '+-- requirements.txt            # Python dependencies\n'
    '+-- .env                        # Environment configuration'
)

# ==================== SECTION 5: DATABASE DESIGN ====================
add_heading('5. Database Design (7 Tables)', level=1)

add_body(
    'We use PostgreSQL with the pgvector extension for storing and querying vector embeddings. '
    'This is critical for semantic search - pgvector lets us do similarity searches directly in SQL.'
)

add_heading('5.1 Entity Relationship Overview', level=2)
add_code(
    'projects (1) ----< (many) documents\n'
    'projects (1) ----< (many) concepts\n'
    'projects (1) ----< (many) brain_state\n'
    'documents (1) ----< (many) chunks       [with 768-dim embedding]\n'
    'documents (1) ----< (many) claims\n'
    'concepts (1) ----< (many) claims\n'
    'concepts (1) ----< (many) relationships (as source)\n'
    'concepts (1) ----< (many) relationships (as target)\n'
    'concepts (1) ----< (many) concepts      [parent-child self-ref]'
)

add_heading('5.2 Key Tables Explained', level=2)

add_body('DOCUMENTS table - Stores uploaded files and their parsed text:')
add_code(
    'documents:\n'
    '  id              UUID (primary key)\n'
    '  project_id      FK -> projects\n'
    '  filename         "research_paper.pdf"\n'
    '  file_path        "uploads/abc123.pdf"\n'
    '  file_type        "pdf" | "docx"\n'
    '  content_text     Full extracted text (can be huge)\n'
    '  metadata_json    {author, title, pages, created_date}'
)

add_body('CHUNKS table - Text segments with vector embeddings:')
add_code(
    'chunks:\n'
    '  id              UUID\n'
    '  document_id     FK -> documents\n'
    '  content          "This paragraph discusses..."\n'
    '  chunk_index      0, 1, 2, 3...  (order in document)\n'
    '  embedding        Vector(768)  <-- pgvector type!\n'
    '  metadata_json    {start_char, end_char}'
)

add_body('CONCEPTS table - Extracted knowledge concepts:')
add_code(
    'concepts:\n'
    '  id                UUID\n'
    '  project_id        FK -> projects\n'
    '  name              "transformer architecture"\n'
    '  description        "Neural network architecture using self-attention..."\n'
    '  generality_score   0.8  (how broad/general, 0-1)\n'
    '  coverage_score     0.6  (how well-covered in documents, 0-1)\n'
    '  consensus_score    0.9  (how much documents agree, 0-1)\n'
    '  embedding          Vector(768)\n'
    '  is_gap             True/False\n'
    '  gap_type           "missing_link" | "under_explored" | "contradictory" | "isolated"\n'
    '  parent_concept_id  FK -> concepts (self-referential hierarchy)\n'
    '  cluster_label      Integer (from HDBSCAN)'
)

add_body('CLAIMS table - Evidence linking documents to concepts:')
add_code(
    'claims:\n'
    '  id              UUID\n'
    '  document_id     FK -> documents\n'
    '  concept_id      FK -> concepts\n'
    '  claim_text       "The authors demonstrate that transformers outperform..."\n'
    '  claim_type       "supports" | "contradicts" | "extends" | "complements"\n'
    '  confidence        0.85 (0-1)\n'
    '  embedding         Vector(768)'
)

add_body('RELATIONSHIPS table - Connections between concepts:')
add_code(
    'relationships:\n'
    '  source_concept_id  FK -> concepts\n'
    '  target_concept_id  FK -> concepts\n'
    '  relationship_type  "prerequisite" | "builds_on" | "contradicts" |\n'
    '                     "complements" | "similar" | "parent_child"\n'
    '  strength           0.75 (0-1)\n'
    '  confidence          0.9 (0-1)\n'
    '  evidence_json       [{claim_id, text, type}]'
)

# ==================== SECTION 6: ML/AI PIPELINE ====================
add_heading('6. ML/AI Pipeline (The Core Value)', level=1)

add_body(
    'This is the heart of Lacuna - the ML pipeline that transforms raw documents into '
    'structured knowledge. Here\'s each stage in detail:'
)

add_heading('6.1 Document Processing Pipeline', level=2)
add_code(
    'PDF/DOCX Upload\n'
    '     |\n'
    '     v\n'
    '[1. PARSE] PyMuPDF extracts text from PDF\n'
    '           python-docx extracts from DOCX\n'
    '           Tesseract OCR as fallback for scanned pages\n'
    '     |\n'
    '     v\n'
    '[2. CHUNK] Split into ~500 token segments\n'
    '           50 token overlap for context preservation\n'
    '           Semantic splitting (respects paragraph boundaries)\n'
    '     |\n'
    '     v\n'
    '[3. EMBED] Each chunk -> 768-dim vector via Ollama\n'
    '           Model: nomic-embed-text\n'
    '           Batch processing (10 at a time)\n'
    '           Stored in pgvector for similarity search'
)

add_heading('6.2 Knowledge Extraction Pipeline', level=2)
add_code(
    'All Document Texts\n'
    '     |\n'
    '     v\n'
    '[4. EXTRACT CONCEPTS] LLM (qwen2.5:3b via Ollama) reads each doc\n'
    '                      Extracts: concept name, description, generality score\n'
    '                      Structured JSON output\n'
    '     |\n'
    '     v\n'
    '[5. EXTRACT CLAIMS]   LLM finds claims supporting/contradicting each concept\n'
    '                      Assigns claim_type and confidence\n'
    '     |\n'
    '     v\n'
    '[6. NORMALIZE]        Deduplicates concepts:\n'
    '                      - Exact name matching (case-insensitive)\n'
    '                      - Embedding similarity > 0.85 = same concept\n'
    '                      - Merges descriptions, averages scores\n'
    '     |\n'
    '     v\n'
    '[7. CLUSTER]          HDBSCAN groups related concepts\n'
    '                      No need to specify K (unlike K-means)\n'
    '                      Handles outliers naturally\n'
    '                      Detects hierarchy via generality scores\n'
    '     |\n'
    '     v\n'
    '[8. RELATIONSHIPS]    Detects connections between concepts:\n'
    '                      - Embedding similarity\n'
    '                      - Generality score differences\n'
    '                      - Claim type analysis\n'
    '                      Types: prerequisite, builds_on, contradicts,\n'
    '                             complements, similar, parent_child\n'
    '     |\n'
    '     v\n'
    '[9. GAP DETECTION]    4 types of gaps:\n'
    '                      - Missing links (similar but unconnected)\n'
    '                      - Under-explored (low coverage score < 0.3)\n'
    '                      - Contradictory (conflicting claims, consensus < 0.5)\n'
    '                      - Isolated (few connections to other concepts)'
)

add_heading('6.3 RAG (Retrieval-Augmented Generation) Pipeline', level=2)
add_code(
    'User Query: "What do the papers say about attention mechanisms?"\n'
    '     |\n'
    '     v\n'
    '[1. EMBED QUERY]      Query -> 768-dim vector\n'
    '     |\n'
    '     v\n'
    '[2. SEMANTIC SEARCH]  Find top-K similar chunks (cosine similarity)\n'
    '                      Find top-K similar concepts\n'
    '                      Uses pgvector for fast vector search\n'
    '     |\n'
    '     v\n'
    '[3. BUILD CONTEXT]    Combine relevant chunks + concept info\n'
    '                      into a structured prompt\n'
    '     |\n'
    '     v\n'
    '[4. LLM GENERATE]     qwen2.5:3b generates answer using context\n'
    '                      Grounded in actual document content\n'
    '     |\n'
    '     v\n'
    '[5. RESPONSE]         Answer + source documents + related concepts\n'
    '                      + confidence score'
)

# ==================== SECTION 7: API ENDPOINTS ====================
add_heading('7. All API Endpoints', level=1)

add_heading('Health', level=2)
add_code('GET  /api/health              # System status (DB, Ollama, version)')

add_heading('Documents', level=2)
add_code(
    'POST /api/documents/upload     # Upload PDF/DOCX -> parse -> chunk -> embed\n'
    'GET  /api/documents/           # List all documents\n'
    'GET  /api/documents/{id}       # Get document details + chunk count\n'
    'DELETE /api/documents/{id}     # Delete document (cascades to chunks, claims)'
)

add_heading('Concepts', level=2)
add_code(
    'POST /api/concepts/extract     # Background: LLM extracts concepts from all docs\n'
    'GET  /api/concepts/            # List concepts (optional: filter gaps only)\n'
    'GET  /api/concepts/map         # Full concept map (nodes + edges + clusters)\n'
    'POST /api/concepts/cluster     # Background: HDBSCAN clustering + hierarchy\n'
    'POST /api/concepts/detect-gaps # Background: Run gap detection algorithm'
)

add_heading('Brain (RAG)', level=2)
add_code(
    'POST /api/brain/query          # Ask a question -> semantic search + LLM answer\n'
    'GET  /api/brain/consensus      # Get current consensus state\n'
    'POST /api/brain/build-consensus # Calculate consensus from all claims'
)

# ==================== SECTION 8: TECH STACK ====================
add_heading('8. Complete Technology Stack', level=1)

table3 = doc.add_table(rows=13, cols=3)
table3.style = 'Light Grid Accent 1'
table3.rows[0].cells[0].text = 'Category'
table3.rows[0].cells[1].text = 'Technology'
table3.rows[0].cells[2].text = 'Why We Chose It'
for cell in table3.rows[0].cells:
    for p in cell.paragraphs:
        p.runs[0].bold = True

stack = [
    ('Web Framework', 'FastAPI 0.115', 'Async, auto-docs, type validation, fastest Python framework'),
    ('Server', 'Uvicorn (ASGI)', 'Async server for FastAPI, production-grade'),
    ('Database', 'PostgreSQL 16', 'Robust, supports pgvector extension'),
    ('Vector Store', 'pgvector', 'Vector similarity search inside PostgreSQL (no separate DB needed)'),
    ('ORM', 'SQLAlchemy 2.0 (async)', 'Industry standard Python ORM, async support'),
    ('Migrations', 'Alembic', 'Database schema versioning (like Prisma Migrate)'),
    ('LLM', 'Ollama + qwen2.5:3b', 'Local LLM, no API costs, privacy-preserving'),
    ('Embeddings', 'Ollama + nomic-embed-text', 'Local 768-dim embeddings, good quality'),
    ('Clustering', 'HDBSCAN', 'Density-based, no K needed, handles outliers'),
    ('Doc Parsing', 'PyMuPDF + python-docx', 'Fast PDF parsing, DOCX support'),
    ('OCR', 'Tesseract (pytesseract)', 'Fallback for scanned/image PDFs'),
    ('Validation', 'Pydantic v2', 'Request/response validation (like Zod for Python)'),
]
for i, (cat, tech, why) in enumerate(stack, 1):
    table3.rows[i].cells[0].text = cat
    table3.rows[i].cells[1].text = tech
    table3.rows[i].cells[2].text = why

# ==================== SECTION 9: QUESTIONS FOR MENTOR ====================
add_heading('9. Questions & Discussion Points for Mentor', level=1)

add_heading('ML Pipeline Questions:', level=2)
add_bullet('Is qwen2.5:3b sufficient for concept/claim extraction, or should we use a larger model?')
add_bullet('HDBSCAN parameters (min_cluster_size=3, min_samples=2) - how to tune for research papers?')
add_bullet('Embedding similarity threshold 0.85 for dedup - is this too aggressive/conservative?')
add_bullet('Should we fine-tune the embedding model on academic text?')
add_bullet('Gap detection heuristics vs ML-based approach - tradeoffs?')
add_bullet('Consensus scoring algorithm - is simple claim-type aggregation enough?')

add_heading('Architecture Questions:', level=2)
add_bullet('Background tasks vs proper job queue (Celery/Redis) for long extraction tasks?')
add_bullet('Is pgvector sufficient or should we use a dedicated vector DB (Pinecone, Qdrant)?')
add_bullet('Chunking strategy: 500 tokens with 50 overlap - optimal for academic papers?')
add_bullet('Should we add a caching layer for embedding queries?')
add_bullet('Multi-user architecture: how to add auth without rewriting everything?')

add_heading('Scaling Questions:', level=2)
add_bullet('Current design is single-project. How to scale to many users/projects?')
add_bullet('Ollama is local-only. When to switch to API-based LLM (OpenAI, Anthropic)?')
add_bullet('How to handle very large document collections (1000+ papers)?')
add_bullet('Batch processing strategy for initial concept extraction?')

add_heading('ML-Specific Improvements:', level=2)
add_bullet('Should we add citation graph analysis?')
add_bullet('Would knowledge graph approaches (Neo4j) be better than relational DB for relationships?')
add_bullet('Active learning: can user feedback improve extraction quality over time?')
add_bullet('Should we add topic modeling (LDA/BERTopic) alongside concept extraction?')

# ==================== SECTION 10: CURRENT LIMITATIONS ====================
add_heading('10. Current Limitations & Known Issues', level=1)

add_bullet('No authentication (single-user, hardcoded project ID)', bold_prefix='Auth:')
add_bullet('Background tasks have no progress tracking or error recovery', bold_prefix='Jobs:')
add_bullet('No caching for repeated embedding queries', bold_prefix='Performance:')
add_bullet('LLM extraction is slow for large documents (sequential processing)', bold_prefix='Speed:')
add_bullet('No unit tests or integration tests yet', bold_prefix='Testing:')
add_bullet('File uploads not scanned for malicious content', bold_prefix='Security:')
add_bullet('No retry logic for Ollama failures', bold_prefix='Resilience:')

# ==================== SECTION 11: WHAT I NEED HELP WITH ====================
add_heading('11. What I Need Guidance On', level=1)

add_body(
    'I\'m a frontend developer (Next.js) learning backend and ML integration. '
    'I need advice on:'
)
add_bullet('Is this architecture sound for what we\'re building?', bold_prefix='1.')
add_bullet('What ML improvements would have the highest impact?', bold_prefix='2.')
add_bullet('How should I evaluate extraction quality (metrics, benchmarks)?', bold_prefix='3.')
add_bullet('What should I prioritize: better models, better algorithms, or better infrastructure?', bold_prefix='4.')
add_bullet('How to properly test an ML pipeline like this?', bold_prefix='5.')
add_bullet('Production readiness: what\'s the gap between current state and deployable?', bold_prefix='6.')

# Save
doc.save('Lacuna_Backend_Architecture.docx')
print("Document saved as Lacuna_Backend_Architecture.docx")
